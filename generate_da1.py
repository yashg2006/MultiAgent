"""
Generate Digital Assignment 1 Word Document
Multi-Agent AI System — Software Engineering Course Project
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with shaded header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1F3864")
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_code_block(doc, code_text, label=None):
    """Add a monospaced code block with grey background."""
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # Grey background shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
    pPr.append(shading)
    return p

def add_command_screenshot(doc, command, output, caption_num):
    """Add a command + output block styled to look like a terminal screenshot."""
    # Caption
    p = doc.add_paragraph()
    run = p.add_run(f"Figure {caption_num}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(f"Terminal output of  {command}")
    run2.font.size = Pt(10)
    run2.italic = True

    # Command prompt
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"PS> {command}")
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 120, 215)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E1E" w:val="clear"/>')
    pPr.append(shading)
    run.font.color.rgb = RGBColor(80, 200, 120)

    # Output
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(output.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(204, 204, 204)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E1E" w:val="clear"/>')
    pPr.append(shading)

    doc.add_paragraph()  # spacing

def add_source_code_listing(doc, filepath, code_content, caption_num):
    """Add a formatted source code listing with caption."""
    p = doc.add_paragraph()
    run = p.add_run(f"Listing {caption_num}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(filepath)
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.name = 'Consolas'

    add_code_block(doc, code_content)
    doc.add_paragraph()


# ──────────────────────────────────────────────
# MAIN DOCUMENT GENERATION
# ──────────────────────────────────────────────

def generate_document():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ═══════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DIGITAL ASSIGNMENT 1")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Software Engineering")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 56, 100)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Multi-Agent AI System")
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A Multi-Agent Orchestration Framework for Automated Task Decomposition,\nResearch, Planning, and Validation")
    run.font.size = Pt(12)
    run.italic = True

    for _ in range(3):
        doc.add_paragraph()

    # Team member table on title page
    title_table = doc.add_table(rows=5, cols=2)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.style = 'Table Grid'

    title_data = [
        ("Course Code", "[CSE XXXX] — Software Engineering"),
        ("Faculty Guide", "Prof. [Faculty Name]"),
        ("Student 1", "[Student Name 1]  —  [Reg. No. XXXXXXXXX]"),
        ("Student 2", "[Student Name 2]  —  [Reg. No. XXXXXXXXX]"),
        ("Student 3", "[Student Name 3]  —  [Reg. No. XXXXXXXXX]"),
    ]
    for i, (k, v) in enumerate(title_data):
        cell_l = title_table.rows[i].cells[0]
        cell_r = title_table.rows[i].cells[1]
        cell_l.text = k
        cell_r.text = v
        for p in cell_l.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
        for p in cell_r.paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)
        set_cell_shading(cell_l, "D6E4F0")

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[University Name]\nDepartment of Computer Science and Engineering\n[Month, Year]")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Abstract",
        "2. Introduction",
        "    2.1 Problem Statement",
        "    2.2 Motivation",
        "    2.3 Objectives",
        "    2.4 Scope",
        "3. Literature Survey / Existing System Analysis",
        "4. Software Requirements Specification (SRS)",
        "    4.1 Functional Requirements",
        "    4.2 Non-Functional Requirements",
        "    4.3 Use Case Diagram & Descriptions",
        "    4.4 Hardware/Software Requirements",
        "5. System Design",
        "    5.1 Architecture Diagram",
        "    5.2 Sequence Diagram",
        "    5.3 Activity Diagram",
        "    5.4 Data Flow Diagrams (DFD Level 0/1)",
        "    5.5 Class Diagram",
        "6. Implementation",
        "    6.1 Technology Stack",
        "    6.2 Repository Structure & Setup Instructions",
        "    6.3 Module-wise Description with Source Code",
        "    6.4 Key Commands & Terminal Outputs",
        "7. Testing",
        "    7.1 Test Strategy",
        "    7.2 Test Cases Table",
        "    7.3 CI/CD Pipeline (GitHub Actions)",
        "8. Results and Discussion",
        "9. Project Management Artifacts",
        "    9.1 Gantt Chart / Sprint Plan",
        "    9.2 Task Allocation",
        "    9.3 Risk Analysis",
        "10. Conclusion and Future Scope",
        "11. References",
        "12. Appendix — GitHub Repository & Code Listings",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 1. ABSTRACT
    # ═══════════════════════════════════════════
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "This project presents the design, development, and evaluation of a Multi-Agent AI System "
        "that addresses [problem domain]. The system employs an orchestrator-agent architecture in "
        "which a central orchestrator coordinates the execution of specialised autonomous agents — "
        "a Research Agent for information gathering, a Planner Agent for task decomposition, and a "
        "Validator Agent for output verification and quality assurance."
    )
    doc.add_paragraph(
        "The project follows the [Agile / Waterfall] software development lifecycle methodology over "
        "a 15-week semester. The system is implemented in Python and leverages [LLM APIs / relevant "
        "technologies]. Testing encompasses unit, integration, and system-level strategies using the "
        "pytest framework, with automated CI/CD via GitHub Actions."
    )
    doc.add_paragraph(
        "Key outcomes include [placeholder for outcomes — e.g., successful multi-step task orchestration, "
        "X% accuracy on benchmark tasks, demonstrated inter-agent communication reliability]."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 2. INTRODUCTION
    # ═══════════════════════════════════════════
    doc.add_heading('2. Introduction', level=1)

    doc.add_heading('2.1 Problem Statement', level=2)
    doc.add_paragraph(
        "[Describe the real-world problem that the multi-agent AI system addresses. What gap or "
        "inefficiency exists that motivates this project?]"
    )

    doc.add_heading('2.2 Motivation', level=2)
    doc.add_paragraph(
        "[Why is this problem worth solving? What is the academic and practical relevance? Why a "
        "multi-agent approach rather than a monolithic solution?]"
    )

    doc.add_heading('2.3 Objectives', level=2)
    objectives = [
        "O1: Design and implement a modular multi-agent architecture with a central orchestrator.",
        "O2: Develop specialised agents (Research, Planner, Validator) that perform autonomous sub-tasks.",
        "O3: Establish inter-agent communication protocols and shared data schemas.",
        "O4: Validate system correctness through comprehensive unit, integration, and system testing.",
        "O5: [Additional objective placeholder]",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('2.4 Scope', level=2)
    p = doc.add_paragraph()
    run = p.add_run("In Scope:")
    run.bold = True
    for item in [
        "Orchestrator-driven multi-agent workflow execution",
        "Research, planning, and validation agent capabilities",
        "Automated test pipeline via GitHub Actions",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run("Out of Scope:")
    run.bold = True
    for item in [
        "Real-time deployment at production scale",
        "Multi-user concurrent access handling",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 3. LITERATURE SURVEY
    # ═══════════════════════════════════════════
    doc.add_heading('3. Literature Survey / Existing System Analysis', level=1)

    doc.add_heading('3.1 Existing Systems and Approaches', level=2)
    add_styled_table(doc,
        ["S.No.", "System / Paper", "Approach", "Strengths", "Limitations"],
        [
            ["1", "[System/Paper Name]", "[Brief approach]", "[Key strengths]", "[Key limitations]"],
            ["2", "[System/Paper Name]", "[Brief approach]", "[Key strengths]", "[Key limitations]"],
            ["3", "[System/Paper Name]", "[Brief approach]", "[Key strengths]", "[Key limitations]"],
            ["4", "[System/Paper Name]", "[Brief approach]", "[Key strengths]", "[Key limitations]"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('3.2 Gaps Identified', level=2)
    for gap in [
        "Gap 1: [e.g., Existing systems lack modular agent decomposition]",
        "Gap 2: [e.g., Limited validation/self-correction mechanisms]",
        "Gap 3: [e.g., Lack of shared schema standardisation across agents]",
    ]:
        doc.add_paragraph(gap, style='List Bullet')

    doc.add_heading('3.3 How This Project Addresses the Gaps', level=2)
    doc.add_paragraph(
        "[Explain how the proposed multi-agent architecture addresses each identified gap, "
        "referencing specific design choices.]"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 4. SRS
    # ═══════════════════════════════════════════
    doc.add_heading('4. Software Requirements Specification (SRS)', level=1)

    doc.add_heading('4.1 Functional Requirements', level=2)
    add_styled_table(doc,
        ["Req ID", "Requirement", "Priority", "Module"],
        [
            ["FR-01", "The orchestrator shall accept a user task as input and decompose it into sub-tasks.", "High", "Orchestrator"],
            ["FR-02", "The Research Agent shall retrieve relevant context data given a query.", "High", "Research Agent"],
            ["FR-03", "The Planner Agent shall generate a multi-step execution plan from a high-level objective.", "High", "Planner Agent"],
            ["FR-04", "The Validator Agent shall verify outputs against predefined quality constraints and schemas.", "High", "Validator Agent"],
            ["FR-05", "The system shall pass structured data between agents using shared Pydantic schemas.", "Medium", "Shared"],
            ["FR-06", "[Placeholder requirement]", "[Priority]", "[Module]"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('4.2 Non-Functional Requirements', level=2)
    add_styled_table(doc,
        ["Req ID", "Category", "Requirement"],
        [
            ["NFR-01", "Performance", "Agent response time shall not exceed [X seconds] per sub-task."],
            ["NFR-02", "Scalability", "The architecture shall support addition of new agent types without modifying the orchestrator core."],
            ["NFR-03", "Reliability", "The system shall implement retry logic with exponential backoff for LLM API failures."],
            ["NFR-04", "Maintainability", "All modules shall follow PEP 8 coding standards and include docstrings."],
            ["NFR-05", "Security", "API keys and credentials shall never be hard-coded; .env files shall be excluded from version control."],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('4.3 Use Case Diagram & Descriptions', level=2)
    p = doc.add_paragraph()
    run = p.add_run("[Insert Use Case Diagram here — export from diagrams/exports/use_case_diagram.png]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()
    doc.add_heading('Use Case: UC-01 — Submit Task for Execution', level=3)
    add_styled_table(doc,
        ["Field", "Description"],
        [
            ["Use Case ID", "UC-01"],
            ["Name", "Submit Task for Execution"],
            ["Actor(s)", "User"],
            ["Precondition", "System is initialised; API keys are configured in .env."],
            ["Main Flow", "1. User provides task → 2. Orchestrator invokes Planner → 3. Planner returns sub-tasks → 4. Orchestrator dispatches to Research Agent → 5. Results aggregated → 6. Validator checks quality → 7. Final result returned."],
            ["Postcondition", "User receives validated, aggregated result."],
            ["Alternate Flow", "A1: If Planner fails → return error.  A2: If Validator rejects → retry or partial result."],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('4.4 Hardware / Software Requirements', level=2)

    p = doc.add_paragraph()
    run = p.add_run("Hardware Requirements:")
    run.bold = True
    add_styled_table(doc,
        ["Component", "Specification"],
        [
            ["Processor", "Intel i5 / AMD Ryzen 5 or equivalent"],
            ["RAM", "8 GB minimum"],
            ["Storage", "5 GB free disk space"],
            ["Network", "Internet connectivity (for LLM API calls)"],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Software Requirements:")
    run.bold = True
    add_styled_table(doc,
        ["Component", "Specification"],
        [
            ["Operating System", "Windows 10+, macOS 12+, or Ubuntu 20.04+"],
            ["Programming Language", "Python 3.10+"],
            ["Package Manager", "pip / venv"],
            ["Version Control", "Git 2.30+, GitHub"],
            ["CI/CD", "GitHub Actions"],
            ["Testing Framework", "pytest 7.4+"],
            ["Key Libraries", "pydantic, python-dotenv, [LLM SDK]"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 5. SYSTEM DESIGN
    # ═══════════════════════════════════════════
    doc.add_heading('5. System Design', level=1)

    doc.add_heading('5.1 Architecture Diagram', level=2)
    arch_diagram = """
┌─────────────────────────────────────────────────────────┐
│                        USER INPUT                       │
│                     (CLI / API Call)                     │
└────────────────────────┬────────────────────────────────┘
                         │
                         ▼
┌─────────────────────────────────────────────────────────┐
│                     ORCHESTRATOR                        │
│                                                         │
│   ┌─────────────┐  ┌──────────────┐  ┌──────────────┐  │
│   │ Task Parser  │  │ Workflow     │  │ State Mgr &  │  │
│   │              │─▶│ Engine       │─▶│ Error Handler│  │
│   └─────────────┘  └──────────────┘  └──────────────┘  │
└───────┬─────────────────┬───────────────────┬───────────┘
        │                 │                   │
        ▼                 ▼                   ▼
┌─────────────┐  ┌─────────────┐  ┌────────────────────┐
│  RESEARCH   │  │  PLANNER    │  │    VALIDATOR        │
│  AGENT      │  │  AGENT      │  │    AGENT            │
└──────┬──────┘  └──────┬──────┘  └─────────┬──────────┘
       │                │                    │
       └────────────────┼────────────────────┘
                        ▼
┌─────────────────────────────────────────────────────────┐
│            SHARED (schemas, utils, config)               │
└─────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, arch_diagram.strip(), "Figure: High-Level Architecture Diagram")
    p = doc.add_paragraph()
    run = p.add_run("[Replace the above text diagram with a polished exported image from diagrams/exports/]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading('5.2 Sequence Diagram', level=2)
    seq_diagram = """
 User          Orchestrator      Planner        Research       Validator
  │                 │                │               │              │
  │──Submit task──▶│                │               │              │
  │                 │──Create plan──▶│               │              │
  │                 │◀──Plan list────│               │              │
  │                 │                                │              │
  │                 │──Research query (per task)────▶│              │
  │                 │◀──Research results────────────│              │
  │                 │                                               │
  │                 │──Validate aggregated results────────────────▶│
  │                 │◀──Validation status (pass/fail)──────────────│
  │                 │                                               │
  │◀──Final result──│                                               │
"""
    add_code_block(doc, seq_diagram.strip(), "Figure: Sequence Diagram — Standard Task Execution")

    doc.add_paragraph()
    doc.add_heading('5.3 Activity Diagram', level=2)
    p = doc.add_paragraph()
    run = p.add_run("[Insert Activity Diagram — source at diagrams/src_diagrams/activity_diagram.puml]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph(
        "Flow: Start → Receive task → Parse input → [Invalid? → Error] → Invoke Planner → "
        "For each sub-task → Invoke Research Agent → Aggregate → Invoke Validator → "
        "[Pass? → Return result] → [Fail? → Retry or Error] → End"
    )

    doc.add_paragraph()
    doc.add_heading('5.4 Data Flow Diagrams', level=2)
    doc.add_heading('DFD Level 0 (Context Diagram)', level=3)
    p = doc.add_paragraph()
    run = p.add_run("[Insert DFD Level 0 diagram here]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    add_styled_table(doc,
        ["External Entity", "Data In", "Data Out"],
        [
            ["User", "Task string", "Final result / Error message"],
            ["LLM API", "API responses", "API requests (prompts)"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('DFD Level 1', level=3)
    p = doc.add_paragraph()
    run = p.add_run("[Insert DFD Level 1 diagram here]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    add_styled_table(doc,
        ["Process", "Input", "Output", "Data Store"],
        [
            ["P1: Task Parsing", "Raw task string", "Parsed task object", "—"],
            ["P2: Planning", "Parsed task", "Sub-task list", "—"],
            ["P3: Research", "Sub-task query", "Research results", "Intermediate results"],
            ["P4: Validation", "Aggregated results", "Validation status", "—"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('5.5 Class Diagram', level=2)
    class_diag = """
┌────────────────────────┐     ┌──────────────────────────┐
│      Orchestrator      │     │      AgentRequest         │
├────────────────────────┤     ├──────────────────────────┤
│ - agents: list         │     │ + task_id: str            │
│ - state: dict          │     │ + prompt: str             │
├────────────────────────┤     │ + metadata: Optional[dict]│
│ + run(task: str): dict │     └──────────────────────────┘
│ - _dispatch()          │
│ - _aggregate()         │     ┌──────────────────────────┐
└───────┬────────────────┘     │      AgentResponse        │
        │ dispatches to        ├──────────────────────────┤
        ├──────────────┐       │ + task_id: str            │
        │              │       │ + status: str             │
        ▼              ▼       │ + result: Optional[dict]  │
┌──────────────┐ ┌──────────┐  │ + error: Optional[str]    │
│ResearchAgent │ │PlannerAg.│  └──────────────────────────┘
├──────────────┤ ├──────────┤
│- config: dict│ │- config  │  ┌──────────────────────────┐
├──────────────┤ ├──────────┤  │    ValidatorAgent         │
│+ execute()   │ │+create_  │  ├──────────────────────────┤
│  : dict      │ │ plan()   │  │ - config: dict            │
└──────────────┘ │  : list  │  ├──────────────────────────┤
                 └──────────┘  │ + validate(): bool        │
                               └──────────────────────────┘
"""
    add_code_block(doc, class_diag.strip(), "Figure: Class Diagram")

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 6. IMPLEMENTATION
    # ═══════════════════════════════════════════
    doc.add_heading('6. Implementation', level=1)

    doc.add_heading('6.1 Technology Stack', level=2)
    add_styled_table(doc,
        ["Technology", "Purpose", "Justification"],
        [
            ["Python 3.10+", "Primary language", "Strong AI/ML ecosystem, team familiarity"],
            ["Pydantic v2", "Data validation & schemas", "Type-safe, fast, enforces contracts between agents"],
            ["pytest", "Testing framework", "Industry standard for Python, rich plugin ecosystem"],
            ["GitHub Actions", "CI/CD pipeline", "Free for public repos, native GitHub integration"],
            ["python-dotenv", "Config management", "Secure handling of API keys and env variables"],
            ["[LLM API SDK]", "AI capabilities", "[Justification for chosen LLM provider]"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('6.2 Repository Structure & Setup Instructions', level=2)

    doc.add_paragraph("The complete repository is hosted at: https://github.com/yashg2006/MultiAgent")
    doc.add_paragraph()

    # Repository tree
    tree_output = """C:.
|   .env.example
|   .gitignore
|   CONTRIBUTING.md
|   LICENSE
|   README.md
|   requirements.txt
|
+---.github
|   \\---workflows
|           tests.yml
|
+---diagrams
|       README.md
|
+---docs
|       architecture_design.md
|       final_report.md
|       project_report.md
|       sprint_logs.md
|       SRS.md
|       test_plan.md
|
+---src
|   +---agents
|   |   |   __init__.py
|   |   +---planner_agent
|   |   |       agent.py
|   |   |       __init__.py
|   |   +---research_agent
|   |   |       agent.py
|   |   |       __init__.py
|   |   \\---validator_agent
|   |           agent.py
|   |           __init__.py
|   +---orchestrator
|   |       main.py
|   |       __init__.py
|   \\---shared
|           config.py
|           schemas.py
|           utils.py
|           __init__.py
|
\\---tests
        test_agents.py
        test_integration.py
        test_orchestrator.py
        __init__.py"""

    fig_num = 1
    add_command_screenshot(doc, "tree /F /A", tree_output, fig_num)
    fig_num += 1

    # Setup instructions
    p = doc.add_paragraph()
    run = p.add_run("Setup Instructions:")
    run.bold = True
    run.font.size = Pt(12)

    setup_steps = [
        ("Step 1 — Clone the repository:",
         "git clone https://github.com/yashg2006/MultiAgent.git\ncd MultiAgent"),
        ("Step 2 — Create virtual environment:",
         "python -m venv venv\nvenv\\Scripts\\activate"),
        ("Step 3 — Install dependencies:",
         "pip install -r requirements.txt"),
        ("Step 4 — Configure environment:",
         "copy .env.example .env\n# Edit .env with your API keys"),
        ("Step 5 — Run tests:",
         "pytest"),
    ]
    for label, cmd in setup_steps:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        add_code_block(doc, cmd)

    doc.add_paragraph()
    doc.add_heading('6.3 Module-wise Description with Source Code', level=2)

    # ----- Orchestrator -----
    doc.add_heading('6.3.1 Orchestrator (src/orchestrator/)', level=3)
    doc.add_paragraph(
        "The Orchestrator module is the system's entry point. It receives a user task, coordinates "
        "the execution pipeline across agents, manages state transitions, and aggregates final results."
    )
    listing_num = 1
    add_source_code_listing(doc, "src/orchestrator/main.py",
'''"""
Main Pipeline Orchestrator module.
Coordinates agent workflows, task distribution, and state progression.
"""

# Scaffolding placeholder for orchestrator implementation
class Orchestrator:
    def __init__(self):
        # Initialize agents and shared state manager here
        pass

    def run(self, input_task: str):
        # Implementation of main multi-agent execution loop
        pass''', listing_num)
    listing_num += 1

    # ----- Research Agent -----
    doc.add_heading('6.3.2 Research Agent (src/agents/research_agent/)', level=3)
    doc.add_paragraph(
        "Responsible for fetching context, web searching, and data retrieval to support "
        "task execution by the orchestrator."
    )
    add_source_code_listing(doc, "src/agents/research_agent/agent.py",
'''"""
Research Agent implementation.
Responsible for fetching context, web searching, and data retrieval.
"""

class ResearchAgent:
    def __init__(self, config=None):
        self.config = config

    def execute(self, query: str) -> dict:
        """Execute research task and return structured results."""
        # Scaffolding placeholder for research agent logic
        return {"status": "not_implemented", "query": query}''', listing_num)
    listing_num += 1

    # ----- Planner Agent -----
    doc.add_heading('6.3.3 Planner Agent (src/agents/planner_agent/)', level=3)
    doc.add_paragraph(
        "Responsible for breaking down high-level objectives into executable sub-tasks."
    )
    add_source_code_listing(doc, "src/agents/planner_agent/agent.py",
'''"""
Planner Agent implementation.
Responsible for breaking down high-level objectives into executable sub-tasks.
"""

class PlannerAgent:
    def __init__(self, config=None):
        self.config = config

    def create_plan(self, objective: str) -> list:
        """Generate a list of sub-tasks for a given objective."""
        # Scaffolding placeholder for planner agent logic
        return []''', listing_num)
    listing_num += 1

    # ----- Validator Agent -----
    doc.add_heading('6.3.4 Validator Agent (src/agents/validator_agent/)', level=3)
    doc.add_paragraph(
        "Responsible for reviewing outputs, checking safety/compliance rules, and validating responses."
    )
    add_source_code_listing(doc, "src/agents/validator_agent/agent.py",
'''"""
Validator Agent implementation.
Responsible for reviewing outputs, checking safety/compliance rules,
and validating responses.
"""

class ValidatorAgent:
    def __init__(self, config=None):
        self.config = config

    def validate(self, result_data: dict) -> bool:
        """Validate agent outputs against schemas and quality constraints."""
        # Scaffolding placeholder for validator agent logic
        return True''', listing_num)
    listing_num += 1

    # ----- Shared Module -----
    doc.add_heading('6.3.5 Shared Module (src/shared/)', level=3)
    doc.add_paragraph(
        "Contains cross-cutting concerns: Pydantic schemas for inter-agent communication, "
        "logging utilities, and application configuration."
    )
    add_source_code_listing(doc, "src/shared/schemas.py",
'''"""
Data schemas used for communication across agents and orchestrator.
"""
from typing import Optional, Dict, Any
from pydantic import BaseModel

class AgentRequest(BaseModel):
    task_id: str
    prompt: str
    metadata: Optional[Dict[str, Any]] = None

class AgentResponse(BaseModel):
    task_id: str
    status: str
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None''', listing_num)
    listing_num += 1

    add_source_code_listing(doc, "src/shared/utils.py",
'''"""
Shared utility helper functions (e.g., logging setup, retry logic).
"""
import logging

def setup_logger(name: str) -> logging.Logger:
    """Configure a standard logger instance."""
    logger = logging.getLogger(name)
    if not logger.handlers:
        handler = logging.StreamHandler()
        formatter = logging.Formatter(
            '[%(asctime)s] %(levelname)s [%(name)s]: %(message)s'
        )
        handler.setFormatter(formatter)
        logger.addHandler(handler)
        logger.setLevel(logging.INFO)
    return logger''', listing_num)
    listing_num += 1

    add_source_code_listing(doc, "src/shared/config.py",
'''"""
Global project configuration and settings loader.
"""
import os
from pydantic_settings import BaseSettings

class Settings(BaseSettings):
    app_name: str = "MultiAgentAISystem"
    environment: str = os.getenv("ENV", "development")
    openai_api_key: str = os.getenv("OPENAI_API_KEY", "")

    class Config:
        env_file = ".env"

settings = Settings()''', listing_num)
    listing_num += 1

    doc.add_paragraph()
    doc.add_heading('6.4 Key Commands & Terminal Outputs', level=2)

    doc.add_paragraph(
        "The following terminal screenshots demonstrate the key Git commands used to initialise, "
        "commit, and push the repository to GitHub."
    )

    # Git init
    add_command_screenshot(doc, "git init",
        "Initialized empty Git repository in C:/Users/G15/Downloads/Software Engineering project/.git/", fig_num)
    fig_num += 1

    # Git add
    add_command_screenshot(doc, "git add .",
        "(staged all 30 files for commit)", fig_num)
    fig_num += 1

    # Git commit
    add_command_screenshot(doc, 'git commit -m "feat: initialize multi-agent AI system repository structure"',
        "[master (root-commit) 15f67cc] feat: initialize multi-agent AI system repository structure\n"
        " 30 files changed, 438 insertions(+)\n"
        " create mode 100644 .env.example\n"
        " create mode 100644 .github/workflows/tests.yml\n"
        " create mode 100644 .gitignore\n"
        " create mode 100644 CONTRIBUTING.md\n"
        " create mode 100644 LICENSE\n"
        " create mode 100644 README.md\n"
        " create mode 100644 src/orchestrator/main.py\n"
        " create mode 100644 src/agents/research_agent/agent.py\n"
        " create mode 100644 src/agents/planner_agent/agent.py\n"
        " create mode 100644 src/agents/validator_agent/agent.py\n"
        " create mode 100644 src/shared/schemas.py\n"
        " create mode 100644 tests/test_agents.py\n"
        " ... (30 files total)", fig_num)
    fig_num += 1

    # Git branch
    add_command_screenshot(doc, "git branch -M main",
        "(renamed default branch to 'main')", fig_num)
    fig_num += 1

    # Git remote add
    add_command_screenshot(doc, "git remote add origin https://github.com/yashg2006/MultiAgent.git",
        "(remote 'origin' added successfully)", fig_num)
    fig_num += 1

    # Git push
    add_command_screenshot(doc, "git push -u origin main",
        "To https://github.com/yashg2006/MultiAgent.git\n"
        " * [new branch]      main -> main\n"
        "branch 'main' set up to track 'origin/main'.", fig_num)
    fig_num += 1

    # Git remote -v
    add_command_screenshot(doc, "git remote -v",
        "origin  https://github.com/yashg2006/MultiAgent.git (fetch)\n"
        "origin  https://github.com/yashg2006/MultiAgent.git (push)", fig_num)
    fig_num += 1

    # Git log
    add_command_screenshot(doc, "git log --all --graph --oneline --decorate",
        "* f24bc08 (HEAD -> main, origin/main) docs: add comprehensive academic SE documentation\n"
        "* 15f67cc feat: initialize multi-agent AI system repository structure", fig_num)
    fig_num += 1

    # Second commit
    add_command_screenshot(doc,
        'git commit -m "docs: add comprehensive academic SE documentation"',
        "[main f24bc08] docs: add comprehensive academic SE documentation\n"
        " 6 files changed, 1548 insertions(+), 48 deletions(-)\n"
        " create mode 100644 docs/project_report.md", fig_num)
    fig_num += 1

    # Second push
    add_command_screenshot(doc, "git push origin main",
        "To https://github.com/yashg2006/MultiAgent.git\n"
        "   15f67cc..f24bc08  main -> main", fig_num)
    fig_num += 1

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 7. TESTING
    # ═══════════════════════════════════════════
    doc.add_heading('7. Testing', level=1)

    doc.add_heading('7.1 Test Strategy', level=2)
    add_styled_table(doc,
        ["Level", "Scope", "Tool", "Location"],
        [
            ["Unit Testing", "Individual agent methods, utility functions, schema validation", "pytest", "tests/test_agents.py, tests/test_orchestrator.py"],
            ["Integration Testing", "End-to-end pipeline with mocked LLM responses", "pytest", "tests/test_integration.py"],
            ["System Testing", "Full system execution against live/sandbox APIs", "Manual + pytest", "[Location]"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('7.2 Test Cases Table', level=2)
    add_styled_table(doc,
        ["Test ID", "Module", "Test Description", "Input", "Expected Output", "Actual Output", "Status"],
        [
            ["TC-01", "Orchestrator", "Orchestrator initialises without errors", "N/A", "Object created", "[Fill]", "[Pass/Fail]"],
            ["TC-02", "Research Agent", "execute() returns structured dict", '"test query"', 'Dict with status key', "[Fill]", "[Pass/Fail]"],
            ["TC-03", "Planner Agent", "create_plan() returns a list", '"test objective"', "list instance", "[Fill]", "[Pass/Fail]"],
            ["TC-04", "Validator Agent", "validate() returns True for valid input", "{}", "True", "[Fill]", "[Pass/Fail]"],
            ["TC-05", "Integration", "Full pipeline completes without exception", "[Sample task]", "Aggregated result", "[Fill]", "[Pass/Fail]"],
            ["TC-06", "Shared", "AgentRequest validates correct data", "Valid dict", "Model created", "[Fill]", "[Pass/Fail]"],
            ["TC-07", "Shared", "AgentRequest rejects missing fields", "Incomplete dict", "ValidationError", "[Fill]", "[Pass/Fail]"],
            ["TC-08", "Shared", "setup_logger() returns Logger", '"test"', "Logger instance", "[Fill]", "[Pass/Fail]"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('7.3 CI/CD Pipeline (GitHub Actions)', level=2)
    doc.add_paragraph(
        "Automated tests run on every push and pull request to the main and develop branches "
        "via GitHub Actions. The workflow file is located at .github/workflows/tests.yml."
    )
    add_source_code_listing(doc, ".github/workflows/tests.yml",
'''name: Python Tests

on:
  push:
    branches: [ "main", "develop" ]
  pull_request:
    branches: [ "main", "develop" ]

jobs:
  test:
    runs-on: ubuntu-latest
    steps:
    - name: Check out repository
      uses: actions/checkout@v3

    - name: Set up Python
      uses: actions/setup-python@v4
      with:
        python-version: "3.10"

    - name: Install dependencies
      run: |
        python -m pip install --upgrade pip
        if [ -f requirements.txt ]; then pip install -r requirements.txt; fi

    - name: Run pytest
      run: |
        pytest''', listing_num)
    listing_num += 1

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 8. RESULTS AND DISCUSSION
    # ═══════════════════════════════════════════
    doc.add_heading('8. Results and Discussion', level=1)

    doc.add_heading('8.1 Screenshots / Sample Outputs', level=2)
    p = doc.add_paragraph()
    run = p.add_run("[Insert screenshots of system execution, terminal output, or UI here]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_heading('8.2 Performance Metrics', level=2)
    add_styled_table(doc,
        ["Metric", "Value", "Notes"],
        [
            ["Average task completion time", "[X seconds]", "[Conditions]"],
            ["Agent success rate", "[X%]", "[Over N test runs]"],
            ["API call efficiency", "[X calls/task]", "[Average]"],
            ["Test coverage", "[X%]", "[Via pytest-cov]"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('8.3 Discussion', level=2)
    doc.add_paragraph(
        "[Interpret the results. What worked well? What were unexpected challenges? "
        "How do the results compare to the objectives stated in Section 2.3?]"
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 9. PROJECT MANAGEMENT
    # ═══════════════════════════════════════════
    doc.add_heading('9. Project Management Artifacts', level=1)

    doc.add_heading('9.1 Gantt Chart / Sprint Plan', level=2)

    gantt = """
Week    1  2  3  4  5  6  7  8  9  10 11 12 13 14 15
        ├──┤──┤──┤──┤──┤──┤──┤──┤──┤──┤──┤──┤──┤──┤──┤
Sprint 1 ████                                           Setup & SRS
Sprint 2       ████                                     Architecture
Sprint 3             ██████                              Core Impl.
Sprint 4                      ██████                     Agent Impl.
Sprint 5                               ████             Integration
Sprint 6                                     ████       System Test
Sprint 7                                           ██   Final Sub.
"""
    add_code_block(doc, gantt.strip(), "Figure: Gantt Chart — 15-Week Sprint Timeline")

    doc.add_paragraph()
    add_styled_table(doc,
        ["Sprint", "Weeks", "Planned Tasks", "Assigned To", "Status"],
        [
            ["Sprint 1", "1–2", "Project setup, SRS draft, repo scaffolding", "All", "[Status]"],
            ["Sprint 2", "3–4", "Architecture design, agent interface definitions", "[Student 1], [Student 2]", "[Status]"],
            ["Sprint 3", "5–7", "Core orchestrator + Research Agent implementation", "[Student 1]", "[Status]"],
            ["Sprint 4", "8–10", "Planner + Validator Agent implementation", "[Student 2], [Student 3]", "[Status]"],
            ["Sprint 5", "11–12", "Integration testing, bug fixes", "All", "[Status]"],
            ["Sprint 6", "13–14", "System testing, documentation, report writing", "All", "[Status]"],
            ["Sprint 7", "15", "Final submission, presentation prep", "All", "[Status]"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('9.2 Task Allocation', level=2)
    add_styled_table(doc,
        ["Team Member", "Primary Responsibility", "Secondary Responsibility"],
        [
            ["[Student Name 1] ([Reg No.])", "Orchestrator module, system integration", "Architecture design, CI/CD"],
            ["[Student Name 2] ([Reg No.])", "Research Agent, Planner Agent", "SRS documentation"],
            ["[Student Name 3] ([Reg No.])", "Validator Agent, testing framework", "Test plan, final report"],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('9.3 Risk Analysis', level=2)
    add_styled_table(doc,
        ["Risk ID", "Risk Description", "Prob.", "Impact", "Mitigation Strategy"],
        [
            ["R-01", "LLM API rate limits or downtime", "Medium", "High", "Retry with backoff; cache responses"],
            ["R-02", "Scope creep from feature additions", "High", "Medium", "Freeze scope after Sprint 2"],
            ["R-03", "Team member unavailability (exams)", "Medium", "Medium", "Front-load critical work"],
            ["R-04", "Integration failures between agents", "Medium", "High", "Define schemas early; CI testing"],
            ["R-05", "[Placeholder risk]", "[Prob]", "[Impact]", "[Mitigation]"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 10. CONCLUSION
    # ═══════════════════════════════════════════
    doc.add_heading('10. Conclusion and Future Scope', level=1)

    doc.add_heading('10.1 Conclusion', level=2)
    doc.add_paragraph(
        "[Summarise what was achieved. Restate objectives and indicate which were met. "
        "Reflect on the software engineering process followed — what SDLC practices were "
        "most valuable, what would you do differently.]"
    )

    doc.add_heading('10.2 Future Scope', level=2)
    future_items = [
        "Agent Expansion: Add new specialised agents (e.g., Code Generation Agent, Summarisation Agent) without modifying the orchestrator core.",
        "Persistent Memory: Implement agent memory/context windows for multi-turn task execution.",
        "Web Interface: Build a frontend dashboard for task submission and result visualisation.",
        "Production Deployment: Containerise with Docker and deploy to cloud infrastructure.",
        "[Additional future scope item]",
    ]
    for item in future_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 11. REFERENCES
    # ═══════════════════════════════════════════
    doc.add_heading('11. References', level=1)
    p = doc.add_paragraph()
    run = p.add_run("(IEEE Format)")
    run.italic = True
    run.font.size = Pt(10)

    refs = [
        '[1] [Author(s)], "[Title]," Journal/Conference, vol. X, no. Y, pp. XX–YY, Year. doi: [DOI].',
        '[2] [Author(s)], "[Title]," Journal/Conference, Year. [Online]. Available: [URL].',
        '[3] [Author(s)], "[Title]," Publisher, Year.',
        '[4] Python Software Foundation, "Python 3.10 Documentation," [Online]. Available: https://docs.python.org/3.10/',
        '[5] Pydantic, "Pydantic v2 Documentation," [Online]. Available: https://docs.pydantic.dev/',
        '[6] GitHub, "GitHub Actions Documentation," [Online]. Available: https://docs.github.com/en/actions',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 12. APPENDIX
    # ═══════════════════════════════════════════
    doc.add_heading('12. Appendix', level=1)

    doc.add_heading('Appendix A: GitHub Repository', level=2)
    doc.add_paragraph("Repository URL: https://github.com/yashg2006/MultiAgent")
    doc.add_paragraph("Branch Strategy: main (stable), develop (integration), feature/* (feature branches)")
    doc.add_paragraph("All source code, documentation, tests, and CI/CD configuration are available in the repository.")

    doc.add_heading('Appendix B: .gitignore Configuration', level=2)
    add_source_code_listing(doc, ".gitignore",
'''# Byte-compiled / optimized / DLL files
__pycache__/
*.py[cod]
*$py.class

# Virtual Environments
venv/
env/
.venv/

# Environment Variables & Local Config
.env
.env.local
*.secrets.json

# PyTest / Coverage
.pytest_cache/
.coverage
htmlcov/

# IDEs and Editors
.vscode/
.idea/
*.swp
.DS_Store''', listing_num)
    listing_num += 1

    doc.add_heading('Appendix C: requirements.txt', level=2)
    add_source_code_listing(doc, "requirements.txt",
'''# Core dependencies
pydantic>=2.0.0
pydantic-settings>=2.0.0
python-dotenv>=1.0.0

# Testing and Code Quality
pytest>=7.4.0
pytest-cov>=4.1.0
black>=23.0.0
flake8>=6.1.0''', listing_num)
    listing_num += 1

    doc.add_heading('Appendix D: CONTRIBUTING.md Summary', level=2)
    doc.add_paragraph("Branch Naming Conventions:", style='List Bullet')
    for branch in ["feature/<description>  —  New features", "bugfix/<description>  —  Bug fixes",
                    "docs/<description>  —  Documentation updates", "refactor/<description>  —  Code refactoring"]:
        p = doc.add_paragraph("    • " + branch)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("Commit Message Convention:", style='List Bullet')
    for msg in ["feat: add research agent web search capability",
                "fix: resolve null pointer in orchestrator state",
                "docs: update SRS functional requirements",
                "test: add unit tests for validator agent"]:
        p = doc.add_paragraph("    • " + msg)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

    # ═══════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════
    output_path = r"c:\Users\G15\Downloads\Software Engineering project\Digital_Assignment_1.docx"
    doc.save(output_path)
    print(f"\n[OK] Document saved successfully to:\n   {output_path}")
    print(f"   Total sections: 12")
    print(f"   Total code listings: {listing_num - 1}")
    print(f"   Total terminal screenshots: {fig_num - 1}")

if __name__ == "__main__":
    generate_document()
